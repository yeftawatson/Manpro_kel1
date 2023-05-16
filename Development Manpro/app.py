from flask import Flask,render_template, request, redirect, url_for, session, flash, jsonify
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import CountVectorizer
import hashlib
import numpy as np
from docx import Document
from flask_mysqldb import MySQL
from uuid import uuid4
import MySQLdb.cursors
import re
import cv2
import io
import nltk
import imghdr
import pdfplumber

nltk.download('wordnet')

app = Flask(__name__, static_folder="D:/Kuliah/Development Manpro/static")

app.secret_key="hehehayoapa"

app.config['MYSQL_HOST'] = 'localhost'
app.config['MYSQL_USER'] = 'root'
app.config['MYSQL_PASSWORD'] = 'proyekmanpro'
app.config['MYSQL_DB'] = 'dbmanpro'
 
mysql = MySQL(app)

allowedExt={'docx', 'pdf', 'jpg'}
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowedExt
@app.route('/hello')
def hello():
    return 'Hello World'

@app.route('/form')
def form():
    return render_template('loginManpro.html')

@app.route('/contact')
def contact():
    return render_template('Contact.html')

@app.route('/aboutus')
def aboutus():
    return render_template('AboutUs.html')
 
@app.route('/signup', methods = ['GET','POST'])
def signup():
    if request.method == 'POST' and 'signup-username' in request.form and 'signup-email' in request.form and 'signup-password' in request.form:
        user_id = str(uuid4())
        username = request.form['signup-username']
        email = request.form['signup-email']
        password = request.form['signup-password']
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute("SELECT * FROM user WHERE user_email = %s", (email,))
        account= cursor.fetchone()
        if account:
            flash('Account already exists')
        elif not re.match(r'[^@]+@[^@]+\.[^@]+',email):
            flash('Invalid email address')
        elif not re.match(r'^[a-zA-Z0-9]+$',username):
            flash('Username must contain only characters and numbers')
        elif not username or not password or not email:
            flash('Please fill out the form')
        else:
            cursor.execute("""INSERT INTO user VALUES(%s,%s,%s,%s)""",(user_id,username,email,password,))
            mysql.connection.commit()
            msg="You have successfully signed up"
        cursor.close()
        return render_template('loginManpro.html', msg=msg)
    return redirect('/form')
@app.route('/login', methods = ['GET','POST'])
def login():
    msg = ''
    if request.method == 'POST' and 'signin-email' in request.form and 'signin-password' in request.form:
        email = request.form['signin-email']
        password = request.form['signin-password']
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute("SELECT * FROM user WHERE user_email = %s and user_pass = %s", (email, password,))
        account = cursor.fetchone()
        if account:
            session['loggedin'] = True
            session['user_id'] = account['user_id']
            session['user_name'] = account['user_name']
            flash('You have successfully logged in')
            redirect ('/form')
        else:
            error = "incorrect password/username"
            return render_template('loginManpro.html', error = error)
    return render_template('loginManpro.html',session=session, msg=msg)
@app.route('/test')
def test():
    return session.get('user_id', 'no user id')
@app.route('/logout')
def logout():
    session.pop('loggedin', None)
    session.pop('user_id', None)
    session.pop('user_name', None)
    return(redirect('/form'))

@app.route('/upload', methods=['GET','POST'])
def upload_file():
    file = request.files['file']
    filename=file.filename
    docid=str(uuid4())
    if file and allowed_file(filename):
        docContent=file.read()
        if 'user_id' in session:
            cursor = mysql.connection.cursor()
            cursor.execute("INSERT INTO document (document_id, user_id, doc_name, doc_content)VALUES (%s, %s, %s, %s)", (docid, session['user_id'], file.filename, docContent, ))
            mysql.connection.commit()
            flash('File uploaded successfully')
            return(render_template('loginManpro.html'))
        else:
            flash('Please login to upload file')
    else:
        return 'Invalid file type'
def jaccard_similarity(set1, set2):
    intersection = len(set1.intersection(set2))
    union = len(set1.union(set2))
    return intersection / union
def preprocess_sentences(text):
    # Replace \n characters with a space
    text = text.replace('\n', ' ')

    # Remove extra spaces between words
    text = re.sub(' +', ' ', text)

    # Remove trailing punctuation marks
    text = re.sub(r'[^\w\s]+$', '', text)

    # Remove consecutive spaces between words and punctuation marks
    text = re.sub(r'\s*([^\w\s])\s*', r' \1 ', text)

    # Normalize spaces around punctuation
    text = re.sub(r'\s([^\w\s])\s', r'\1 ', text)

    # Remove leading and trailing whitespaces from each sentence
    sentences = [sentence.strip() for sentence in text.split('.')]

    return sentences

@app.route('/jaccard_similarity', methods=['GET','POST'])
def calculate_jaccard_similarity():
    # file1 = request.files['tes.docx']
    # file2 = request.files['tes2.docx']
    # doc_id1 = '9c798b43-1d9e-40db-9e7b-2d6144d81ded'
    # doc_id2 = 'd1512997-78ad-4b91-a296-3afc57b35c19'
    doc_id = '9c798b43-1d9e-40db-9e7b-2d6144d81ded'

    cursor = mysql.connection.cursor()
    cursor.execute("SELECT doc_content FROM document WHERE document_id = %s", (doc_id,))
    row = cursor.fetchone()

    if row:
        content1 = row[0]
        if imghdr.what(None, h=content1) is not None:
            return 'Invalid document type'
        if content1.startswith(b'%PDF'):
            pdf = pdfplumber.open(io.BytesIO(content1))
            extracted_text = ""
            for page in pdf.pages:
                extracted_text += page.extract_text()
            paragraf_docx = extracted_text.lower()
        else:
            paragraf_docx = ""
            document = Document(io.BytesIO(content1))
            for paragraph in document.paragraphs:
                paragraf_docx += paragraph.text.lower()
            print(paragraf_docx)

        cursor.execute("SELECT document_id, doc_content FROM document WHERE document_id != %s", (doc_id,))
        similarity_results = []
        for row in cursor.fetchall():
            content2 = row[1]
            if imghdr.what(None, h=content2) is not None:
                continue
            if content2.startswith(b'%PDF'):
                pdf = pdfplumber.open(io.BytesIO(content2))
                extracted_text = ""
                for page in pdf.pages:
                    extracted_text += page.extract_text()
                paragraf_docx2 = extracted_text.lower()
                print(paragraf_docx2)
            else:
                paragraf_docx2 = ""
                document = Document(io.BytesIO(content2))
                for paragraph in document.paragraphs:
                    paragraf_docx2 += paragraph.text.lower()
            sentences1 = preprocess_sentences(paragraf_docx)
            sentences2 = preprocess_sentences(paragraf_docx2)
            arrteks=[]
            arrteks2=[]
            arrResHash = set()
            for sentence in sentences1:
                arrteks.append(sentence)
                sentence_hash = hashlib.sha256(sentence.encode()).hexdigest()
                arrResHash.add(sentence_hash)
            
            arrResHash2 = set()
            for sentence in sentences2:
                arrteks2.append(sentence)
                sentence_hash = hashlib.sha256(sentence.encode()).hexdigest()
                arrResHash2.add(sentence_hash)

            result_jaccard = jaccard_similarity(arrResHash, arrResHash2) * 100
            jaccard_dec = np.around(result_jaccard, decimals=4)
            final_result = '{:.4f}%'.format(jaccard_dec)
            
            similarity_results.append(f"Jaccard similarity with document ID {row[0]}: {final_result}")
            print("arr1")
            print(sentences1)
            print("arr2")
            print (sentences2)
        return "<br>".join(similarity_results)
    else:
        return 'Invalid document ID'
    # cursor = mysql.connection.cursor()
    # cursor.execute("SELECT doc_content FROM document WHERE document_id = %s", (doc_id1,))
    # row1 = cursor.fetchone()
    # cursor.execute("SELECT doc_content FROM document WHERE document_id = %s", (doc_id2,))
    # row2 = cursor.fetchone()
    # if row1 and row2:
    #     content1 = row1[0]
    #     content2 = row2[0]

    #     # Perform the Jaccard similarity calculation
    #     paragraf_docx1 = ""
    #     arrResHash1 = []
    #     for paragraph in Document(io.BytesIO(content1)).paragraphs:
    #         paragraf_docx1 += paragraph.text.lower()
    #         arrTeks = []
    #         arrResHash1 = []
    #         for i in sent_tokenize(paragraf_docx1):
    #             i = i.replace(" ", "")
    #             arrTeks.append(i)
    #             i = hashlib.sha256(i.encode())
    #             i = i.hexdigest()
    #             arrResHash1.append(i)

    #     paragraf_docx2 = ""
    #     arrResHash2 = []
    #     for paragraph in Document(io.BytesIO(content2)).paragraphs:
    #         paragraf_docx2 += paragraph.text.lower()
    #         arrTeks2 = []
    #         arrResHash2 = []
    #         for i in sent_tokenize(paragraf_docx2):
    #             i = i.replace(" ", "")
    #             arrTeks2.append(i)
    #             i = hashlib.sha256(i.encode())
    #             i = i.hexdigest()
    #             arrResHash2.append(i)

    #     result_jaccard = jaccard_similarity(arrResHash1, arrResHash2) * 100
    #     jaccard_dec = np.around(result_jaccard, decimals=4)
    #     final_result = '{:.4f}%'.format(jaccard_dec)

    #     return jsonify({'result': final_result, 'arr1': arrResHash1, 'arr2': arrResHash2})
    # else:
    #     return 'Invalid document IDs'
    # document = Document('D:/MATERI KULIAH SUTAN/SEMESTER 6/MANPRO/tes.docx')
    # document2 = Document('D:/MATERI KULIAH SUTAN/SEMESTER 6/MANPRO/tes2.docx')

    # paragraf_docx = ""
    # paragraf_docx2 = ""
    # for paragraph in document.paragraphs:
    #     paragraf_docx += paragraph.text.lower()
   
    # for paraf in document2.paragraphs:
    #     paragraf_docx2 += paraf.text.lower()

    # arrTeks = []
    # arrResHash = []
    # for i in sent_tokenize(paragraf_docx):
    #     i = i.replace(" ", "")
    #     arrTeks.append(i)
    #     i = hashlib.sha256(i.encode())
    #     i = i.hexdigest()
    #     arrResHash.append(i)

    # arrTeks2 = []
    # arrResHash2 = []
    # for i in sent_tokenize(paragraf_docx2):
    #     i = i.replace(" ", "")
    #     arrTeks2.append(i)
    #     i = hashlib.sha256(i.encode())
    #     i = i.hexdigest()
    #     arrResHash2.append(i)

    # result_jaccard = jaccard_similarity(arrResHash, arrResHash2) * 100
    # jaccard_dec = np.around(result_jaccard, decimals=4)
    # # final_result = np.char.mod('%.4f%%', jaccard_dec)

    # final_result = '{:.4f}%'.format(jaccard_dec)
    # return jsonify({'result': final_result})

@app.route('/compare_images', methods=['GET','POST'])
def compare_images():
    doc_id='c72f6dc3-21d7-4092-bae0-f84e333e5c5e'
    # Retrieve the image blobs from the database
    cursor = mysql.connection.cursor()
    cursor.execute("SELECT doc_content,doc_name FROM document WHERE document_id!=%s", (doc_id,))
    rows = cursor.fetchall()

    results = []

    # Load the image to compare
    cursor.execute("SELECT doc_content FROM document WHERE document_id=%s", (doc_id,))
    img_to_compare_blob = cursor.fetchone()[0]

    # Convert the blob data to a NumPy array
    nparr = np.frombuffer(img_to_compare_blob, np.uint8)

    # Decode the NumPy array as an image
    img_to_compare = cv2.imdecode(nparr, cv2.IMREAD_COLOR)

    # Convert the image to grayscale
    gray_img_to_compare = cv2.cvtColor(img_to_compare, cv2.COLOR_BGR2GRAY)

    # Calculate the histogram of the image to compare
    hist_to_compare = cv2.calcHist([gray_img_to_compare], [0], None, [256], [0, 256])
    cv2.normalize(hist_to_compare, hist_to_compare, alpha=0, beta=1, norm_type=cv2.NORM_MINMAX)

    # Loop through each row in the database
    for row in rows:
        blob_data = row[0]

        if imghdr.what(None, h=blob_data) is not None:
            # Convert the blob data to a NumPy array
            nparr = np.frombuffer(blob_data, np.uint8)

            # Decode the NumPy array as an image
            dataset_img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)

            # Convert the dataset image to grayscale
            gray_dataset_img = cv2.cvtColor(dataset_img, cv2.COLOR_BGR2GRAY)

            # Calculate the histogram of the dataset image
            hist_dataset_img = cv2.calcHist([gray_dataset_img], [0], None, [256], [0, 256])
            cv2.normalize(hist_dataset_img, hist_dataset_img, alpha=0, beta=1, norm_type=cv2.NORM_MINMAX)

            # Compare the histograms of the two images
            score = cv2.compareHist(hist_to_compare, hist_dataset_img, cv2.HISTCMP_CORREL)

            # Convert the score to a percentage
            percentage_score = score * 100

            # Append the percentage score to the results list
            results.append(f"Image Similarity with Image name {row[1]}: {percentage_score:.2f}")
    # Return the results as a response
    return "<br>".join(results)

app.run(host='localhost', port=5000)


    
    