from nltk.tokenize import sent_tokenize #utk sentence tokenization
from docx import Document #utk baca dokumen
from sklearn.feature_extraction.text import CountVectorizer #utk BoW
import hashlib #utk hashing
import numpy as np

def jaccard_similarity(list1, list2):
    set1 = set(list1)
    set2 = set(list2)
    intersection = set1.intersection(set2)
    union = set1.union(set2)
    return len(intersection) / len(union)

#ambil isi dokumen
document = Document('tes.docx')
document2 = Document('tes2.docx')

paragraf_docx = ""
paragraf_docx2 = ""
for paragraph in document.paragraphs:
    # print(paragraph.text)
    paragraf_docx += paragraph.text.lower()
   
for paraf in document2.paragraphs:
     paragraf_docx2 += paraf.text.lower()

arrTeks = []
arrResHash = []
for i in sent_tokenize(paragraf_docx):
    # print(i)
    i = i.replace(" ", "")
    arrTeks.append(i)
    i = hashlib.sha256(i.encode())
    i = i.hexdigest()
    arrResHash.append(i)

arrTeks2 = []
arrResHash2 = []
for i in sent_tokenize(paragraf_docx2):
    # print(i)
    i = i.replace(" ", "")
    arrTeks2.append(i)
    i = hashlib.sha256(i.encode())
    i = i.hexdigest()
    arrResHash2.append(i)

print("array teks 1", arrTeks)
print("array hash 1", arrResHash)

print("array teks 2", arrTeks2)
print("array hash 2",arrResHash2)

result_jaccard = jaccard_similarity(arrResHash, arrResHash2) * 100
jaccard_dec = np.around(result_jaccard, decimals=4)
final_result = np.char.mod('%.4f%%', jaccard_dec)

print(final_result)

